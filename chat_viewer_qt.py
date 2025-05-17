# PySide6 version of Chat Viewer with modern layout and DOCX export
import sys, os, json
from PySide6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QFileDialog,
                               QLabel, QScrollArea, QFrame, QListWidget, QListWidgetItem)
from PySide6.QtCore import Qt, QSize
from PySide6.QtGui import QTextDocument
from docx import Document
from dateutil import parser
from datetime import datetime

# === Helpers ===
def process_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    parsed = []
    for msg in data.get("messages", []):
        creator = msg.get("creator", {})
        name = creator.get("name", "Deleted User")
        email = creator.get("email", "")
        text = msg.get("text", "")
        ts_raw = msg.get("created_date", "")
        try:
            ts = parser.parse(ts_raw).isoformat()
        except Exception:
            ts = ""
        if text:
            parsed.append({"name": name, "email": email, "text": text, "timestamp": ts})
    return parsed

def format_time(iso):
    try:
        return datetime.fromisoformat(iso).strftime('%I:%M %p').lstrip('0')
    except:
        return ""

def format_date(iso):
    try:
        return datetime.fromisoformat(iso).strftime('%d %b %Y')
    except:
        return "Unknown Date"

def export_docx(messages, title):
    doc = Document()
    doc.add_heading(title, level=1)
    for msg in messages:
        entry = f"{format_time(msg['timestamp'])} - {msg['name']} ({msg['email']})\n{msg['text']}\n"
        doc.add_paragraph(entry)
    path, _ = QFileDialog.getSaveFileName(None, "Save Chat", title + ".docx", "Word Documents (*.docx)")
    if path:
        doc.save(path)

# === UI ===
class ChatViewer(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Chat Viewer - PySide6")
        self.resize(1200, 700)
        self.setStyleSheet("background-color: #1e1e2a; color: white;")

        self.chat_sessions = {}
        self.current_messages = []

        main_layout = QHBoxLayout(self)

        # Sidebar
        sidebar = QVBoxLayout()
        import_btn = QPushButton("+ Import Chat")
        import_btn.clicked.connect(self.import_chat)
        import_btn.setStyleSheet("background-color: #4f93ff; color: white; padding: 10px; font-weight: bold")

        export_btn = QPushButton("Export to DOCX")
        export_btn.clicked.connect(self.export_chat)
        export_btn.setStyleSheet("background-color: #3a3a3c; color: #ddd; padding: 10px")

        self.chat_list = QListWidget()
        self.chat_list.setStyleSheet("background-color: #2b2d3c; color: white")
        self.chat_list.itemClicked.connect(self.load_chat)

        sidebar.addWidget(import_btn)
        sidebar.addWidget(export_btn)
        sidebar.addWidget(self.chat_list)
        sidebar_container = QFrame()
        sidebar_container.setLayout(sidebar)
        sidebar_container.setFixedWidth(240)
        sidebar_container.setStyleSheet("background-color: #2b2d3c")
        main_layout.addWidget(sidebar_container)

        # Main Viewer
        self.chat_area = QVBoxLayout()
        self.chat_area.setAlignment(Qt.AlignTop)

        chat_scroll = QScrollArea()
        chat_scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_content.setLayout(self.chat_area)
        chat_scroll.setWidget(scroll_content)
        chat_scroll.setStyleSheet("background-color: #121218")

        main_layout.addWidget(chat_scroll)

    def import_chat(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Chat JSON", os.getcwd(), "JSON Files (*.json)")
        if path:
            name = os.path.basename(path)
            messages = process_json(path)
            self.chat_sessions[name] = messages
            item = QListWidgetItem(name)
            self.chat_list.addItem(item)

    def export_chat(self):
        current_item = self.chat_list.currentItem()
        if current_item:
            title = current_item.text()
            export_docx(self.chat_sessions[title], title)

    def load_chat(self, item):
        title = item.text()
        self.current_messages = self.chat_sessions.get(title, [])
        self.render_chat()

    def render_chat(self):
        # Clear layout
        for i in reversed(range(self.chat_area.count())):
            widget = self.chat_area.itemAt(i).widget()
            if widget:
                widget.setParent(None)

        last_date = ""
        left_user = self.current_messages[0]['name'] if self.current_messages else ""

        for msg in self.current_messages:
            timestamp = msg.get("timestamp", "")
            date_label = format_date(timestamp)
            if date_label != last_date:
                last_date = date_label
                lbl = QLabel(date_label)
                lbl.setStyleSheet("color: #ccc; background-color: #2b2d3c; padding: 6px; border-radius: 6px")
                self.chat_area.addWidget(lbl)

            is_self = msg["name"] != left_user
            bubble = QLabel()
            bubble.setText(f"<b>{msg['name']}</b> ({msg.get('email','')})<br>{msg['text']}<br><small>{format_time(timestamp)}</small>")
            bubble.setWordWrap(True)
            bubble.setTextFormat(Qt.RichText)
            bubble.setStyleSheet(f"padding:10px; border-radius:12px; background-color: {'#4f93ff' if is_self else '#3a3a3c'}")
            bubble.setMaximumWidth(800)
            bubble.setAlignment(Qt.AlignLeft if not is_self else Qt.AlignRight)

            wrap = QHBoxLayout()
            wrap.addStretch() if is_self else wrap.addSpacing(10)
            wrap.addWidget(bubble)
            wrap.addSpacing(10) if is_self else wrap.addStretch()
            wrap_widget = QWidget()
            wrap_widget.setLayout(wrap)
            self.chat_area.addWidget(wrap_widget)

# === Launch ===
if __name__ == '__main__':
    app = QApplication(sys.argv)
    viewer = ChatViewer()
    viewer.show()
    sys.exit(app.exec())
