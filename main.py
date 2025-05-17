# === Chat Viewer GUI styled to match reference (adapted for 'created_date') ===
import tkinter as tk
from tkinter import filedialog, Canvas, Scrollbar, Frame, Label, Button, Entry
from PIL import Image, ImageTk
import json
from datetime import datetime
import os
from docx import Document
from dateutil import parser

# === Helpers ===
def format_time(iso):
    try:
        return datetime.fromisoformat(iso).strftime('%I:%M %p').lstrip('0')
    except:
        return iso

def process_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    parsed = []
    for msg in data.get("messages", []):
        creator = msg.get("creator", {})
        name = creator.get("name", "Deleted User")
        email = creator.get("email", "")
        text = msg.get("text", "")
        ts_raw = msg.get("created_date", "")
        try:
            ts = parser.parse(ts_raw).isoformat()
        except Exception:
            ts = ""
        if text:
            parsed.append({"name": name, "email": email, "text": text, "timestamp": ts})
    return parsed

def export_to_docx(messages, title):
    doc = Document()
    doc.add_heading(title, level=1)
    for msg in messages:
        entry = f"{format_time(msg['timestamp'])} - {msg['name']} ({msg['email']})\n{msg['text']}\n"
        doc.add_paragraph(entry)
    file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file:
        doc.save(file)

# === GUI ===
class ChatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Modern Chat Viewer")
        self.root.geometry("1280x800")
        self.root.configure(bg="#1e1e2a")

        self.chat_sessions = {}
        self.current_messages = []
        self.active_title = None

        # Left Sidebar
        self.sidebar = Frame(root, bg="#2b2d3c", width=240)
        self.sidebar.pack(side="left", fill="y")

        Label(self.sidebar, text="Chats", font=("Segoe UI", 14, "bold"), fg="#fff", bg="#2b2d3c").pack(pady=10)
        Button(self.sidebar, text="+ Import Chat", font=("Segoe UI", 10, "bold"), bg="#4f93ff", fg="white", bd=0,
               command=self.load_file).pack(padx=10, pady=(0, 10), fill="x")
        Button(self.sidebar, text="Export to DOCX", font=("Segoe UI", 10), bg="#3a3a3c", fg="#ddd", bd=0,
               command=self.export_docx).pack(padx=10, pady=(0, 20), fill="x")

        self.chat_buttons_frame = Frame(self.sidebar, bg="#2b2d3c")
        self.chat_buttons_frame.pack(fill="both", expand=True)

        # Main Chat Area
        self.chat_main = Frame(root, bg="#121218")
        self.chat_main.pack(side="left", fill="both", expand=True)

        self.chat_header = Frame(self.chat_main, bg="#1e1e2a")
        self.chat_header.pack(fill="x")
        self.chat_title = Label(self.chat_header, text="Select a chat", font=("Segoe UI", 16, "bold"), fg="#f2f2f2", bg="#1e1e2a")
        self.chat_title.pack(side="left", padx=15, pady=10)

        self.canvas = Canvas(self.chat_main, bg="#121218", highlightthickness=0)
        self.scrollbar = Scrollbar(self.chat_main, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.inner_frame = Frame(self.canvas, bg="#121218")
        self.canvas.create_window((0, 0), window=self.inner_frame, anchor="nw")

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.inner_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind_all("<MouseWheel>", lambda e: self.canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))

        # Bottom Input Bar (Placeholder)
        self.input_bar = Frame(self.chat_main, bg="#1e1e2a")
        self.input_bar.pack(fill="x", side="bottom")
        Entry(self.input_bar, font=("Segoe UI", 10), fg="#fff", bg="#2b2d3c", bd=0, insertbackground='white').pack(padx=15, pady=12, fill="x")

    def load_file(self):
        path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if not path:
            return
        display_name = os.path.basename(path)
        base_name = display_name
        counter = 2
        while display_name in self.chat_sessions:
            display_name = f"{base_name} ({counter})"
            counter += 1
        messages = process_json(path)
        self.chat_sessions[display_name] = messages

        btn = Button(self.chat_buttons_frame, text=display_name, bg="#3a3a3c", fg="#fff", font=("Segoe UI", 10), relief="flat",
                     activebackground="#4f93ff", command=lambda m=messages, t=display_name: self.select_chat(m, t))
        btn.pack(fill="x", padx=10, pady=5)
        self.select_chat(messages, display_name)

    def export_docx(self):
        if not self.current_messages:
            return
        export_to_docx(self.current_messages, self.active_title or "Chat Export")

    def select_chat(self, messages, title):
        self.current_messages = messages
        self.active_title = title
        self.chat_title.config(text=title)
        self.render_chat()

    def render_chat(self):
        for widget in self.inner_frame.winfo_children():
            widget.destroy()

        if not self.current_messages:
            Label(self.inner_frame, text="No messages to show", fg="white", bg="#121218").pack()
            return

        left_user = self.current_messages[0]["name"]
        last_date = ""

        for msg in self.current_messages:
            timestamp = msg["timestamp"]
            try:
                dt = datetime.fromisoformat(timestamp)
                current_date = dt.strftime('%d %b %Y')
            except Exception:
                dt = None
                current_date = "Unknown Date"

            if current_date != last_date:
                last_date = current_date
                Label(self.inner_frame, text=current_date, bg="#2b2d3c", fg="#ccc", font=("Segoe UI", 9, "bold"),
                      padx=10, pady=4).pack(pady=(10, 2))

            is_self = msg["name"] != left_user
            bubble_color = "#4f93ff" if is_self else "#3a3a3c"

            bubble = Frame(self.inner_frame, bg=bubble_color, padx=12, pady=8)
            bubble.pack(anchor="e" if is_self else "w", pady=6, padx=20, fill="x")
            Label(bubble, text=msg["name"], font=("Segoe UI", 10, "bold"), fg="#f0f0f0", bg=bubble_color).pack(anchor="w")
            Label(bubble, text=msg.get("email", ""), font=("Segoe UI", 8), fg="#ccc", bg=bubble_color).pack(anchor="w")
            Label(bubble, text=msg["text"], font=("Segoe UI", 10), wraplength=700, justify="left", fg="white", bg=bubble_color).pack(anchor="w", pady=(4, 4))
            if dt:
                Label(bubble, text=dt.strftime('%I:%M %p').lstrip('0'), font=("Segoe UI", 8), fg="#aaa", bg=bubble_color).pack(anchor="e")

# === Launch ===
if __name__ == "__main__":
    root = tk.Tk()
    app = ChatApp(root)
    root.mainloop()
